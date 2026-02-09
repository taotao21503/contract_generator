#!/usr/bin/env python3
"""
签收单批量生成工具
从"原始数据"表读取合同信息，匹配产品明细表，生成签收单文档
"""

import sys
import os
from pathlib import Path

# 添加当前目录到路径
sys.path.insert(0, str(Path(__file__).parent))

# 导入主模块的依赖检查
from contract_generator import ensure_dependencies
if not ensure_dependencies():
    sys.exit(1)

from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
import pandas as pd
import tempfile
import os

from contract_generator import (
    replace_placeholders_in_paragraph,
    replace_placeholders_in_table,
    table_to_image,
    insert_image_at_placeholder,
    insert_table_at_placeholder,
    read_excel_table_from_row,
    find_placeholder_paragraph,
    convert_to_pdf
)


def read_unique_contracts(excel_path: str, sheet_name: str = None, header_row: int = 1) -> list[dict]:
    """
    读取Excel中的唯一合同记录

    Args:
        excel_path: Excel文件路径
        sheet_name: 工作表名称，None则自动检测
        header_row: 表头所在行号（从1开始）

    Returns:
        唯一合同列表
    """
    wb = load_workbook(excel_path, read_only=True)

    # 自动检测工作表
    if sheet_name is None:
        # 优先查找 "原始数据"，否则使用第一个工作表
        if "原始数据" in wb.sheetnames:
            sheet_name = "原始数据"
        else:
            sheet_name = wb.sheetnames[0]
        print(f"  使用工作表: {sheet_name}")

    ws = wb[sheet_name]

    rows = list(ws.iter_rows(values_only=True))
    if not rows or len(rows) < header_row:
        return []

    # 使用指定行作为表头
    header_idx = header_row - 1
    headers = [str(h).strip() if h else f"Column_{i}" for i, h in enumerate(rows[header_idx])]

    # 读取表头之后的数据
    data = []
    seen_contracts = set()

    for row in rows[header_idx + 1:]:
        if not any(row):
            continue
        record = {}
        for i, value in enumerate(row):
            if i < len(headers):
                record[headers[i]] = str(value) if value is not None else ""

        # 按合同编号去重
        contract_no = record.get("合同编号", "").strip()
        if contract_no and contract_no not in seen_contracts:
            seen_contracts.add(contract_no)
            data.append(record)

    wb.close()
    return data


def find_detail_sheet_by_key(detail_excel_path: str, customer: str, contract_no: str) -> str:
    """
    在明细Excel文件中查找对应的工作表

    Args:
        detail_excel_path: 明细Excel文件路径
        customer: 客户名称
        contract_no: 合同编号

    Returns:
        工作表名称，未找到返回None
    """
    target_key = f"{customer}{contract_no}"

    try:
        wb = load_workbook(detail_excel_path, read_only=True)

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            # 读取第一行的key值（在B1单元格）
            first_row = list(ws.iter_rows(max_row=1, values_only=True))
            if first_row and first_row[0]:
                row_values = first_row[0]
                if len(row_values) > 1 and row_values[1]:
                    sheet_key = str(row_values[1]).strip()
                    if sheet_key == target_key:
                        wb.close()
                        return sheet_name

        wb.close()
    except Exception as e:
        print(f"    查找工作表失败: {e}")

    return None


def generate_receipt(
    template_path: str,
    data: dict,
    output_path: str,
    detail_excel_path: str,
    detail_start_row: int = 4,
    use_image: bool = True,
    generate_pdf: bool = False
):
    """
    生成单个签收单

    Args:
        template_path: Word模板路径
        data: 合同数据字典
        output_path: 输出文件路径
        detail_excel_path: 明细Excel文件路径
        detail_start_row: 明细数据起始行号
        use_image: 是否使用图片方式插入明细表格
        generate_pdf: 是否同时生成PDF文件
    """
    doc = Document(template_path)

    # 替换所有段落中的占位符
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, data)

    # 替换所有表格中的占位符
    for table in doc.tables:
        replace_placeholders_in_table(table, data)

    # 替换页眉页脚中的占位符
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                replace_placeholders_in_paragraph(paragraph, data)
            for table in section.header.tables:
                replace_placeholders_in_table(table, data)
        if section.footer:
            for paragraph in section.footer.paragraphs:
                replace_placeholders_in_paragraph(paragraph, data)
            for table in section.footer.tables:
                replace_placeholders_in_table(table, data)

    # 查找明细工作表
    detail_found = False
    customer = data.get("收货方名称（乙方）", "").strip()
    contract_no = data.get("合同编号", "").strip()

    sheet_name = find_detail_sheet_by_key(detail_excel_path, customer, contract_no)

    if sheet_name:
        table_data = read_excel_table_from_row(detail_excel_path, detail_start_row, sheet_name)

        if table_data:
            placeholder = "~~产品明细~~"

            if use_image:
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                    tmp_path = tmp.name
                try:
                    if table_to_image(table_data, tmp_path):
                        insert_image_at_placeholder(doc, placeholder, tmp_path, width_inches=6.0)
                        detail_found = True
                    else:
                        insert_table_at_placeholder(doc, placeholder, table_data)
                        detail_found = True
                finally:
                    if os.path.exists(tmp_path):
                        os.remove(tmp_path)
            else:
                insert_table_at_placeholder(doc, placeholder, table_data)
                detail_found = True

    if not detail_found:
        # 没有找到明细，清除占位符
        paragraph, _ = find_placeholder_paragraph(doc, "~~产品明细~~")
        if paragraph:
            paragraph.clear()
            paragraph.add_run("（产品明细待补充）")

    doc.save(output_path)

    # 生成PDF
    pdf_path = None
    if generate_pdf:
        pdf_path = convert_to_pdf(output_path)

    return detail_found, pdf_path


def batch_generate_receipts(
    data_excel: str,
    detail_excel: str,
    template_path: str,
    output_dir: str,
    header_row: int = 1,
    detail_start_row: int = 4,
    use_image: bool = True,
    generate_pdf: bool = False,
    limit: int = None
):
    """
    批量生成签收单

    Args:
        data_excel: 合同数据Excel路径（包含"原始数据"工作表）
        detail_excel: 产品明细Excel路径
        template_path: Word模板路径
        output_dir: 输出目录
        detail_start_row: 明细数据起始行号
        use_image: 是否使用图片方式
        generate_pdf: 是否生成PDF
        limit: 限制生成数量（用于测试）
    """
    # 确保输出目录存在
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    # 读取合同数据
    contracts = read_unique_contracts(data_excel, header_row=header_row)

    if not contracts:
        print("错误: 未找到合同数据")
        return

    if limit:
        contracts = contracts[:limit]

    print(f"共 {len(contracts)} 个合同待处理")
    print("-" * 50)

    success_count = 0
    fail_count = 0
    pdf_count = 0

    for i, contract in enumerate(contracts, start=1):
        try:
            contract_no = contract.get("合同编号", "未知")
            customer = contract.get("收货方名称（乙方）", "未知客户")

            # 生成文件名
            filename = f"签收单-{contract_no}-{customer}.docx"
            # 清理非法字符
            filename = "".join(c if c not in r'<>:"/\|?*' else "_" for c in filename)

            output_file = output_path / filename

            detail_found, pdf_path = generate_receipt(
                template_path,
                contract,
                str(output_file),
                detail_excel,
                detail_start_row=detail_start_row,
                use_image=use_image,
                generate_pdf=generate_pdf
            )

            success_count += 1
            detail_info = " (含明细)" if detail_found else " (无明细)"
            pdf_info = " +PDF" if pdf_path else ""
            if pdf_path:
                pdf_count += 1
            print(f"[{i}/{len(contracts)}] {contract_no}{detail_info}{pdf_info}")

        except Exception as e:
            fail_count += 1
            print(f"[{i}/{len(contracts)}] 失败: {e}")

    print("-" * 50)
    print(f"完成! 成功: {success_count}, 失败: {fail_count}")
    if generate_pdf:
        print(f"PDF生成: {pdf_count}")


def main():
    import argparse

    parser = argparse.ArgumentParser(description="签收单批量生成工具")
    parser.add_argument("--data", "-d", required=True, help="合同数据Excel文件路径")
    parser.add_argument("--detail", "-D", required=True, help="产品明细Excel文件路径")
    parser.add_argument("--template", "-t", required=True, help="Word模板文件路径")
    parser.add_argument("--output", "-o", default="output/receipts", help="输出目录")
    parser.add_argument("--header-row", type=int, default=1, help="主数据表头所在行号")
    parser.add_argument("--start-row", type=int, default=4, help="明细数据起始行号")
    parser.add_argument("--table", action="store_true", help="使用表格方式而非图片")
    parser.add_argument("--pdf", action="store_true", help="同时生成PDF")
    parser.add_argument("--limit", type=int, help="限制生成数量（用于测试）")

    args = parser.parse_args()

    # 验证文件存在
    for path, name in [(args.data, "数据文件"), (args.detail, "明细文件"), (args.template, "模板文件")]:
        if not Path(path).exists():
            print(f"错误: {name}不存在: {path}")
            return 1

    print(f"数据文件: {args.data}")
    print(f"明细文件: {args.detail}")
    print(f"模板文件: {args.template}")
    print(f"输出目录: {args.output}")
    print(f"数据表头行: {args.header_row}")
    print(f"明细起始行: {args.start_row}")
    print(f"插入方式: {'表格' if args.table else '图片'}")
    if args.limit:
        print(f"限制数量: {args.limit}")
    print("-" * 50)

    batch_generate_receipts(
        args.data,
        args.detail,
        args.template,
        args.output,
        header_row=args.header_row,
        detail_start_row=args.start_row,
        use_image=not args.table,
        generate_pdf=args.pdf,
        limit=args.limit
    )

    return 0


if __name__ == "__main__":
    exit(main())
