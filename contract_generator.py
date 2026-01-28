#!/usr/bin/env python3
"""
合同批量生成工具
根据Excel中的合同信息，批量生成Word合同文件
"""

import argparse
import re
import sys
import subprocess
from pathlib import Path
from typing import Optional


# 依赖包配置: {导入名: (pip包名, 用途说明)}
REQUIRED_PACKAGES = {
    "openpyxl": ("openpyxl", "读取Excel文件"),
    "docx": ("python-docx", "读写Word文档"),
}

# PyPI镜像源配置: (名称, URL)
PIP_MIRRORS = [
    ("官方PyPI (国外)", None),  # None表示使用默认源
    ("清华大学 (推荐)", "https://pypi.tuna.tsinghua.edu.cn/simple"),
    ("阿里云", "https://mirrors.aliyun.com/pypi/simple"),
    ("中国科技大学", "https://pypi.mirrors.ustc.edu.cn/simple"),
    ("华为云", "https://repo.huaweicloud.com/repository/pypi/simple"),
]


def check_python_version():
    """检查Python版本"""
    if sys.version_info < (3, 9):
        print(f"错误: Python版本过低")
        print(f"  当前版本: {sys.version}")
        print(f"  最低要求: Python 3.9+")
        print(f"\n请升级Python后重试")
        return False
    return True


def check_dependencies():
    """
    检查依赖包是否已安装

    Returns:
        (是否全部安装, 缺失的包列表)
    """
    missing = []

    for import_name, (pip_name, description) in REQUIRED_PACKAGES.items():
        try:
            __import__(import_name)
        except ImportError:
            missing.append((import_name, pip_name, description))

    return len(missing) == 0, missing


def select_pip_mirror() -> tuple[str, Optional[str]]:
    """
    让用户选择pip镜像源

    Returns:
        (镜像名称, 镜像URL)，URL为None表示使用默认源
    """
    print("\n请选择下载源:")
    print("-" * 50)
    for i, (name, url) in enumerate(PIP_MIRRORS, start=1):
        print(f"  {i}. {name}")
    print("-" * 50)

    while True:
        try:
            choice = input("请输入编号 [默认2-清华大学]: ").strip()
        except (EOFError, KeyboardInterrupt):
            print("\n已取消")
            return None, None

        # 默认选择清华源
        if choice == "":
            return PIP_MIRRORS[1]

        try:
            index = int(choice) - 1
            if 0 <= index < len(PIP_MIRRORS):
                return PIP_MIRRORS[index]
            else:
                print(f"请输入 1-{len(PIP_MIRRORS)} 之间的数字")
        except ValueError:
            print("请输入有效的数字")


def install_dependencies(packages: list, mirror_url: Optional[str] = None) -> bool:
    """
    安装缺失的依赖包

    Args:
        packages: 缺失的包列表 [(import_name, pip_name, description), ...]
        mirror_url: 镜像源URL，None表示使用默认源

    Returns:
        是否安装成功
    """
    pip_packages = [pip_name for _, pip_name, _ in packages]
    print(f"\n正在安装: {', '.join(pip_packages)}")
    if mirror_url:
        print(f"使用镜像: {mirror_url}")
    print("-" * 50)

    try:
        # 构建pip命令
        cmd = [sys.executable, "-m", "pip", "install"]

        # 如果指定了镜像源，添加参数
        if mirror_url:
            cmd.extend(["-i", mirror_url, "--trusted-host", mirror_url.split("//")[1].split("/")[0]])

        cmd.append("--quiet")
        cmd.extend(pip_packages)

        subprocess.check_call(cmd)
        print("-" * 50)
        print("安装完成!")
        return True
    except subprocess.CalledProcessError as e:
        print(f"\n安装失败: {e}")
        mirror_hint = f" -i {mirror_url}" if mirror_url else ""
        print("\n请手动执行以下命令安装:")
        print(f"  {sys.executable} -m pip install{mirror_hint} {' '.join(pip_packages)}")
        return False
    except FileNotFoundError:
        print("\n错误: 找不到pip")
        print("请确保pip已正确安装，或手动安装依赖:")
        print(f"  pip install {' '.join(pip_packages)}")
        return False


def ensure_dependencies():
    """
    确保所有依赖已安装，如果缺失则提示用户安装

    Returns:
        是否可以继续运行
    """
    # 检查Python版本
    if not check_python_version():
        return False

    # 检查依赖包
    all_installed, missing = check_dependencies()

    if all_installed:
        return True

    # 显示缺失的依赖
    print("检测到缺少以下依赖包:")
    print("-" * 50)
    for import_name, pip_name, description in missing:
        print(f"  - {pip_name}: {description}")
    print("-" * 50)

    # 询问是否自动安装
    while True:
        try:
            response = input("\n是否自动安装缺失的依赖? [Y/n]: ").strip().lower()
        except (EOFError, KeyboardInterrupt):
            print("\n已取消")
            return False

        if response in ("", "y", "yes"):
            # 选择镜像源
            mirror_name, mirror_url = select_pip_mirror()
            if mirror_name is None:
                return False

            print(f"\n已选择: {mirror_name}")

            if install_dependencies(missing, mirror_url):
                # 重新检查是否安装成功
                all_installed, _ = check_dependencies()
                return all_installed
            return False
        elif response in ("n", "no"):
            print("\n请手动安装依赖后重试:")
            pip_packages = [pip_name for _, pip_name, _ in missing]
            print(f"  pip install {' '.join(pip_packages)}")
            print("\n推荐使用国内镜像源加速下载:")
            print(f"  pip install -i https://pypi.tuna.tsinghua.edu.cn/simple {' '.join(pip_packages)}")
            return False
        else:
            print("请输入 Y(是) 或 N(否)")


# 在导入依赖前先检查
if not ensure_dependencies():
    sys.exit(1)

# 依赖检查通过后再导入
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def read_excel_data(excel_path: str, header_row: int = 1) -> list[dict]:
    """
    读取Excel文件，返回合同数据列表

    Args:
        excel_path: Excel文件路径
        header_row: 表头所在行号（从1开始计数，默认为1）

    Returns:
        合同数据列表，每条数据为字典格式
    """
    wb = load_workbook(excel_path, read_only=True)
    ws = wb.active

    rows = list(ws.iter_rows(values_only=True))
    if not rows or len(rows) < header_row:
        return []

    # 指定行作为表头（转换为0索引）
    header_idx = header_row - 1
    headers = [str(h).strip() if h else f"Column_{i}" for i, h in enumerate(rows[header_idx])]

    # 表头之后的行作为数据
    data = []
    for row in rows[header_idx + 1:]:
        # 跳过空行
        if not any(row):
            continue
        record = {}
        for i, value in enumerate(row):
            if i < len(headers):
                # 将值转换为字符串，None转为空字符串
                record[headers[i]] = str(value) if value is not None else ""
        data.append(record)

    wb.close()
    return data


def replace_placeholders_in_paragraph(paragraph, data: dict):
    """
    替换段落中的占位符

    Args:
        paragraph: Word段落对象
        data: 替换数据字典
    """
    # 获取完整文本
    full_text = paragraph.text

    # 检查是否包含占位符
    if "{{" not in full_text:
        return

    # 替换所有占位符
    new_text = full_text
    for key, value in data.items():
        placeholder = "{{" + key + "}}"
        new_text = new_text.replace(placeholder, value)

    # 如果文本有变化，需要更新段落
    if new_text != full_text:
        # 保留第一个run的格式，清空其他run
        if paragraph.runs:
            # 保存第一个run的格式
            first_run = paragraph.runs[0]
            first_run.text = new_text
            # 清空其他run
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            # 如果没有run，直接添加文本
            paragraph.add_run(new_text)


def replace_placeholders_in_table(table, data: dict):
    """
    替换表格中的占位符

    Args:
        table: Word表格对象
        data: 替换数据字典
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_placeholders_in_paragraph(paragraph, data)


def find_detail_excel(data_dir: str, data: dict) -> Optional[str]:
    """
    查找对应的明细Excel文件（单文件模式）

    Args:
        data_dir: 数据目录路径
        data: 合同数据字典

    Returns:
        Excel文件路径，未找到返回None
    """
    contract_no = data.get("合同号", "").strip()
    customer_name = data.get("客户名称", "").strip()
    bu_name = data.get("BU名称", "").strip()

    if not all([contract_no, customer_name, bu_name]):
        return None

    # 尝试多种可能的文件名格式
    possible_names = [
        f"{contract_no}{customer_name}{bu_name}.xlsx",
        f"{contract_no}+{customer_name}+{bu_name}.xlsx",
        f"{contract_no}-{customer_name}-{bu_name}.xlsx",
        f"{contract_no}_{customer_name}_{bu_name}.xlsx",
    ]

    data_path = Path(data_dir)
    for name in possible_names:
        file_path = data_path / name
        if file_path.exists():
            return str(file_path)

    return None


def find_detail_sheet(detail_excel_path: str, data: dict,
                      customer_field: str = "收货方名称（乙方）",
                      contract_field: str = "合同编号") -> Optional[str]:
    """
    在明细Excel文件中查找对应的工作表（多工作表模式）

    Args:
        detail_excel_path: 明细Excel文件路径
        data: 合同数据字典
        customer_field: 客户名称字段名
        contract_field: 合同编号字段名

    Returns:
        工作表名称，未找到返回None
    """
    customer = data.get(customer_field, "").strip()
    contract_no = data.get(contract_field, "").strip()

    if not customer or not contract_no:
        return None

    target_key = f"{customer}{contract_no}"

    try:
        wb = load_workbook(detail_excel_path, read_only=True)

        # 遍历所有工作表，查找key匹配的
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            # 读取第一行的key值（通常在B1单元格）
            first_row = list(ws.iter_rows(max_row=1, values_only=True))
            if first_row and first_row[0]:
                row_values = first_row[0]
                # key可能在第2列（索引1）
                if len(row_values) > 1 and row_values[1]:
                    sheet_key = str(row_values[1]).strip()
                    if sheet_key == target_key:
                        wb.close()
                        return sheet_name

        wb.close()
    except Exception:
        pass

    return None


def read_excel_table_from_row(excel_path: str, start_row: int = 9,
                               sheet_name: Optional[str] = None) -> list[list[str]]:
    """
    读取Excel文件指定行之后的表格数据

    Args:
        excel_path: Excel文件路径
        start_row: 起始行号（从1开始计数，读取该行及之后的数据）
        sheet_name: 工作表名称，None表示使用活动工作表

    Returns:
        表格数据，二维列表
    """
    wb = load_workbook(excel_path, read_only=True)

    if sheet_name:
        ws = wb[sheet_name]
    else:
        ws = wb.active

    table_data = []
    for row_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
        if row_idx >= start_row:
            # 跳过完全空的行
            if any(cell is not None for cell in row):
                row_data = [str(cell) if cell is not None else "" for cell in row]
                table_data.append(row_data)

    wb.close()
    return table_data


def set_cell_border(cell, border_color: str = "000000", border_size: str = "4"):
    """为单元格设置边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), border_size)
        border.set(qn('w:color'), border_color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_cell_shading(cell, color: str):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color)
    tcPr.append(shading)


def set_cell_vertical_alignment(cell, align: str = "center"):
    """设置单元格垂直对齐"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)


def is_number_like(value: str) -> bool:
    """判断字符串是否像数字（用于右对齐）"""
    if not value or value.strip() == "":
        return False
    # 去除常见的数字格式字符
    cleaned = value.replace(",", "").replace(" ", "").replace("￥", "").replace("¥", "")
    try:
        float(cleaned)
        return True
    except ValueError:
        return False


def append_table_to_doc(doc, table_data: list[list[str]], title: str = None):
    """
    将表格数据追加到Word文档末尾

    Args:
        doc: Word文档对象
        table_data: 表格数据，二维列表
        title: 可选的表格标题
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    if not table_data:
        return

    # 添加空行
    doc.add_paragraph()

    # 添加标题（如果有）
    if title:
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(14)
        title_run.font.name = '黑体'
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 计算列数（取最大列数，但限制合理范围）
    max_cols = max(len(row) for row in table_data)
    # 去除尾部空列
    while max_cols > 1:
        all_empty = all(
            len(row) < max_cols or not row[max_cols - 1].strip()
            for row in table_data
        )
        if all_empty:
            max_cols -= 1
        else:
            break

    # 创建表格
    table = doc.add_table(rows=len(table_data), cols=max_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置表格自动调整
    table.autofit = True

    # 填充数据
    for row_idx, row_data in enumerate(table_data):
        row = table.rows[row_idx]
        is_header = row_idx == 0  # 第一行是表头

        for col_idx in range(max_cols):
            cell = row.cells[col_idx]
            cell_value = row_data[col_idx] if col_idx < len(row_data) else ""

            # 清理单元格内容中的换行符
            cell_value = cell_value.replace("\n", " ").strip()
            cell.text = cell_value

            # 设置边框
            set_cell_border(cell)

            # 设置垂直居中
            set_cell_vertical_alignment(cell, "center")

            # 设置段落和字体样式
            for paragraph in cell.paragraphs:
                # 设置段落对齐
                if is_header:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif is_number_like(cell_value):
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                    if is_header:
                        run.bold = True
                        run.font.size = Pt(10)

            # 表头行设置背景色
            if is_header:
                set_cell_shading(cell, "D9E2F3")  # 浅蓝色背景


def generate_contract(template_path: str, data: dict, output_path: str,
                      data_dir: Optional[str] = None,
                      detail_excel_path: Optional[str] = None,
                      detail_start_row: int = 4):
    """
    根据模板生成单个合同文件

    Args:
        template_path: Word模板路径
        data: 合同数据字典
        output_path: 输出文件路径
        data_dir: 数据目录路径，用于查找明细Excel文件（单文件模式）
        detail_excel_path: 明细Excel文件路径（多工作表模式）
        detail_start_row: 明细数据起始行号（默认4）
    """
    doc = Document(template_path)

    # 替换所有段落中的占位符
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, data)

    # 替换所有表格中的占位符
    for table in doc.tables:
        replace_placeholders_in_table(table, data)

    # 替换页眉页脚中的占位符
    for section in doc.sections:
        # 页眉
        if section.header:
            for paragraph in section.header.paragraphs:
                replace_placeholders_in_paragraph(paragraph, data)
            for table in section.header.tables:
                replace_placeholders_in_table(table, data)
        # 页脚
        if section.footer:
            for paragraph in section.footer.paragraphs:
                replace_placeholders_in_paragraph(paragraph, data)
            for table in section.footer.tables:
                replace_placeholders_in_table(table, data)

    # 查找并追加明细表格
    detail_found = False

    # 模式1: 多工作表模式（明细在一个Excel文件的多个工作表中）
    if detail_excel_path and Path(detail_excel_path).exists():
        sheet_name = find_detail_sheet(detail_excel_path, data)
        if sheet_name:
            table_data = read_excel_table_from_row(detail_excel_path, detail_start_row, sheet_name)
            if table_data:
                append_table_to_doc(doc, table_data, title="合同配置清单")
                detail_found = True

    # 模式2: 单文件模式（每个合同对应一个明细Excel文件）
    if not detail_found and data_dir:
        detail_excel = find_detail_excel(data_dir, data)
        if detail_excel:
            table_data = read_excel_table_from_row(detail_excel, start_row=9)
            if table_data:
                append_table_to_doc(doc, table_data, title="合同配置清单")
                detail_found = True

    doc.save(output_path)
    return detail_found


def sanitize_filename(name: str) -> str:
    """
    清理文件名中的非法字符

    Args:
        name: 原始文件名

    Returns:
        清理后的文件名
    """
    # 替换Windows和Unix都不允许的字符
    illegal_chars = r'[<>:"/\\|?*]'
    return re.sub(illegal_chars, "_", name)


def generate_output_filename(data: dict,
                             contract_field: str = "合同编号",
                             customer_field: str = "收货方名称（乙方）",
                             bu_field: str = "BU") -> str:
    """
    根据合同数据生成输出文件名

    Args:
        data: 合同数据字典
        contract_field: 合同编号字段名
        customer_field: 客户名称字段名
        bu_field: BU字段名

    Returns:
        输出文件名（不含路径）
    """
    # 尝试多种可能的字段名
    contract_no = (data.get(contract_field, "").strip() or
                   data.get("合同号", "").strip() or "未知合同号")
    customer_name = (data.get(customer_field, "").strip() or
                     data.get("客户名称", "").strip() or "未知客户")
    bu_name = (data.get(bu_field, "").strip() or
               data.get("BU名称", "").strip() or "未知BU")

    filename = f"{contract_no}-{customer_name}-{bu_name}.docx"
    return sanitize_filename(filename)


def batch_generate_contracts(
    excel_path: str,
    template_path: str,
    output_dir: str,
    data_dir: Optional[str] = None,
    detail_excel_path: Optional[str] = None,
    header_row: int = 1,
    detail_start_row: int = 4
) -> tuple[int, int, list[str]]:
    """
    批量生成合同文件

    Args:
        excel_path: Excel文件路径
        template_path: Word模板路径
        output_dir: 输出目录路径
        data_dir: 数据目录路径，用于查找明细Excel文件（单文件模式）
        detail_excel_path: 明细Excel文件路径（多工作表模式）
        header_row: 数据Excel表头所在行号（默认1）
        detail_start_row: 明细数据起始行号（默认4）

    Returns:
        (成功数量, 失败数量, 错误信息列表)
    """
    # 确保输出目录存在
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    # 如果未指定data_dir，使用excel_path所在目录
    if data_dir is None:
        data_dir = str(Path(excel_path).parent)

    # 读取Excel数据
    contracts = read_excel_data(excel_path, header_row=header_row)

    if not contracts:
        return 0, 0, ["Excel文件中没有数据"]

    success_count = 0
    fail_count = 0
    errors = []

    for i, contract in enumerate(contracts, start=1):
        try:
            filename = generate_output_filename(contract)
            output_file = output_path / filename
            detail_found = generate_contract(
                template_path, contract, str(output_file),
                data_dir=data_dir,
                detail_excel_path=detail_excel_path,
                detail_start_row=detail_start_row
            )
            success_count += 1
            detail_info = " (含明细表格)" if detail_found else ""
            print(f"[{i}/{len(contracts)}] 生成成功: {filename}{detail_info}")
        except Exception as e:
            fail_count += 1
            error_msg = f"第{i}条记录生成失败: {e}"
            errors.append(error_msg)
            print(f"[{i}/{len(contracts)}] {error_msg}")

    return success_count, fail_count, errors


def print_environment_info():
    """打印环境信息"""
    print("环境检查")
    print("=" * 50)
    print(f"Python版本: {sys.version}")
    print(f"Python路径: {sys.executable}")
    print("-" * 50)
    print("依赖包状态:")
    for import_name, (pip_name, description) in REQUIRED_PACKAGES.items():
        try:
            module = __import__(import_name)
            version = getattr(module, "__version__", "未知版本")
            print(f"  [OK] {pip_name} ({version}) - {description}")
        except ImportError:
            print(f"  [缺失] {pip_name} - {description}")
    print("=" * 50)


def main():
    parser = argparse.ArgumentParser(
        description="合同批量生成工具 - 根据Excel中的合同信息批量生成Word合同文件"
    )
    parser.add_argument(
        "--excel", "-e",
        default="data/contracts.xlsx",
        help="合同信息Excel文件路径 (默认: data/contracts.xlsx)"
    )
    parser.add_argument(
        "--template", "-t",
        default="data/template.docx",
        help="Word模板文件路径 (默认: data/template.docx)"
    )
    parser.add_argument(
        "--output", "-o",
        default="output/",
        help="输出目录路径 (默认: output/)"
    )
    parser.add_argument(
        "--detail", "-d",
        default=None,
        help="明细Excel文件路径（多工作表模式）"
    )
    parser.add_argument(
        "--header-row",
        type=int,
        default=1,
        help="数据Excel表头所在行号 (默认: 1)"
    )
    parser.add_argument(
        "--detail-start-row",
        type=int,
        default=4,
        help="明细数据起始行号 (默认: 4)"
    )
    parser.add_argument(
        "--check", "-c",
        action="store_true",
        help="仅检查环境和依赖，不执行生成"
    )

    args = parser.parse_args()

    # 如果是检查模式，打印环境信息后退出
    if args.check:
        print_environment_info()
        return 0

    # 验证输入文件存在
    excel_path = Path(args.excel)
    template_path = Path(args.template)

    if not excel_path.exists():
        print(f"错误: Excel文件不存在: {args.excel}")
        return 1

    if not template_path.exists():
        print(f"错误: 模板文件不存在: {args.template}")
        return 1

    print(f"Excel文件: {args.excel}")
    print(f"模板文件: {args.template}")
    print(f"输出目录: {args.output}")
    if args.detail:
        print(f"明细文件: {args.detail}")
    print(f"表头行号: {args.header_row}")
    print(f"明细起始行: {args.detail_start_row}")
    print("-" * 50)

    success, fail, errors = batch_generate_contracts(
        args.excel,
        args.template,
        args.output,
        detail_excel_path=args.detail,
        header_row=args.header_row,
        detail_start_row=args.detail_start_row
    )

    print("-" * 50)
    print(f"生成完成! 成功: {success}, 失败: {fail}")

    if errors:
        print("\n错误详情:")
        for error in errors:
            print(f"  - {error}")
        return 1

    return 0


if __name__ == "__main__":
    exit(main())
