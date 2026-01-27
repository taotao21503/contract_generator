#!/usr/bin/env python3
"""
创建示例Excel和Word模板文件
"""

from openpyxl import Workbook
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_sample_excel():
    """创建示例Excel文件"""
    wb = Workbook()
    ws = wb.active
    ws.title = "合同信息"

    # 表头
    headers = ["合同号", "客户名称", "BU名称", "签约日期", "合同金额", "联系人", "联系电话"]
    ws.append(headers)

    # 示例数据
    sample_data = [
        ["HT-2024-001", "北京科技有限公司", "华北事业部", "2024-01-15", "100,000.00", "张三", "13800138001"],
        ["HT-2024-002", "上海贸易有限公司", "华东事业部", "2024-01-20", "250,000.00", "李四", "13900139002"],
        ["HT-2024-003", "广州电子有限公司", "华南事业部", "2024-02-01", "180,000.00", "王五", "13700137003"],
    ]

    for row in sample_data:
        ws.append(row)

    # 调整列宽
    column_widths = [15, 25, 15, 15, 15, 10, 15]
    for i, width in enumerate(column_widths, start=1):
        ws.column_dimensions[chr(64 + i)].width = width

    wb.save("data/contracts.xlsx")
    print("示例Excel文件已创建: data/contracts.xlsx")


def create_sample_template():
    """创建示例Word模板"""
    doc = Document()

    # 标题
    title = doc.add_heading("销售合同", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 合同编号
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("合同编号：{{合同号}}")
    run.font.size = Pt(12)

    doc.add_paragraph()

    # 甲乙方信息
    doc.add_paragraph("甲方（卖方）：我方公司名称")
    doc.add_paragraph("乙方（买方）：{{客户名称}}")
    doc.add_paragraph("所属部门：{{BU名称}}")

    doc.add_paragraph()

    # 正文
    doc.add_paragraph("根据《中华人民共和国合同法》及相关法律法规的规定，甲乙双方在平等、自愿、公平、诚实信用的基础上，就产品销售事宜达成如下协议：")

    doc.add_paragraph()

    # 合同条款
    doc.add_heading("一、合同金额", level=1)
    doc.add_paragraph(f"本合同总金额为人民币：{{{{合同金额}}}}元整。")

    doc.add_heading("二、签约日期", level=1)
    doc.add_paragraph(f"本合同签订日期为：{{{{签约日期}}}}。")

    doc.add_heading("三、联系方式", level=1)
    doc.add_paragraph(f"乙方联系人：{{{{联系人}}}}")
    doc.add_paragraph(f"联系电话：{{{{联系电话}}}}")

    doc.add_paragraph()
    doc.add_paragraph()

    # 签章
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "甲方（盖章）："
    table.cell(0, 1).text = "乙方（盖章）："
    table.cell(1, 0).text = "授权代表签字："
    table.cell(1, 1).text = "授权代表签字："
    table.cell(2, 0).text = "日期："
    table.cell(2, 1).text = "日期："

    doc.save("data/template.docx")
    print("示例Word模板已创建: data/template.docx")


if __name__ == "__main__":
    create_sample_excel()
    create_sample_template()
    print("\n示例文件创建完成！")
    print("现在可以运行: python contract_generator.py")
